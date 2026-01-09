
import os
import re
from docx import Document
from docx.shared import Inches, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import RGBColor

# Configuration
REPORT_MD_PATH = r"g:\Drive'ım\Dersler - Projeler\Bitirme Projesi-A Güz Dönemi\reports\final_report.md"
FIGURES_DIR = r"g:\Drive'ım\Dersler - Projeler\Bitirme Projesi-A Güz Dönemi\reports\figures"
OUTPUT_DOCX_PATH = r"g:\Drive'ım\Dersler - Projeler\Bitirme Projesi-A Güz Dönemi\reports\Bitirme_Projesi_Raporu.docx"

STUDENT_ID = "241307109"
DEPT = "Bilişim Sistemleri Mühendisliği 4. Sınıf"
PROJECT_TITLE = "TÜRKİYE İLLERİNİN SOSYO-EKONOMİK GELİŞMİŞLİK DÜZEYLERİNE GÖRE\nMAKİNE ÖĞRENMESİ YÖNTEMLERİYLE KÜMELENMESİ"

def setup_styles(doc):
    # Set default style to Times New Roman 12
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    paragraph_format = style.paragraph_format
    paragraph_format.space_after = Pt(6) # Add space after paragraphs for readability without manual newlines
    
    # Configure headings
    for i in range(1, 4):
        style = doc.styles[f'Heading {i}']
        font = style.font
        font.name = 'Times New Roman'
        font.bold = True
        font.color.rgb = RGBColor(0, 0, 0)
        paragraph_format = style.paragraph_format
        if i == 1:
            font.size = Pt(14)
            font.all_caps = True
            paragraph_format.space_before = Pt(12)
            paragraph_format.space_after = Pt(12)
        else:
            font.size = Pt(12)
            font.all_caps = False
            paragraph_format.space_before = Pt(12)
            paragraph_format.space_after = Pt(6)

    # Configure List Bullet style to ensure font consistency
    if 'List Bullet' in doc.styles:
        style = doc.styles['List Bullet']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)

    # Configure List Number for consistency
    if 'List Number' in doc.styles:
        style = doc.styles['List Number']
        font = style.font
        font.name = 'Times New Roman'
        font.size = Pt(12)
        
    # Configure Table Text style (if we use custom paragraphs in cells)
    # We will enforce runs in add_table_to_doc manually, but good to have base style correct.

def add_formatted_paragraph(doc, text, style=None, alignment=None):
    if style:
        p = doc.add_paragraph(style=style)
    else:
        p = doc.add_paragraph()
        
    if alignment is not None:
        p.alignment = alignment

    # Clean cleanup of common markdown artifacts that shouldn't appear
    # Replace single * if it's acting as bullet in text or weird artifact, 
    # but be careful not to kill mathematical * operations if any exist (though rare in this context)
    # User specifically asked to remove excess * and #.
    
    # Regex to capture bold parts **text**
    # We split by **
    parts = text.split('**')
    
    for i, part in enumerate(parts):
        if not part:
            continue
            
        # Inside the part, we might still have single * italics or other artifacts
        # For this request, we just print the text.
        
        run = p.add_run(part)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        if i % 2 == 1: # Odd indices form the 'bold' part in split
            run.bold = True
            
    return p

def create_cover_page(doc):
    # T.C. Header
    p = doc.add_paragraph('T.C.\nKOCAELİ ÜNİVERSİTESİ\nTEKNOLOJİ FAKÜLTESİ')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.style.font.bold = True
    
    # Logo placeholder or space
    doc.add_paragraph('\n\n\n\n')
    
    # Department
    p = doc.add_paragraph('BİLİŞİM SİSTEMLERİ MÜHENDİSLİĞİ BÖLÜMÜ')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.style.font.bold = True
    p.style.font.size = Pt(14)
    
    doc.add_paragraph('\n\n')
    
    # Title
    p = doc.add_paragraph(PROJECT_TITLE)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.style.font.bold = True
    p.style.font.size = Pt(16)
    
    doc.add_paragraph('\n')
    
    p = doc.add_paragraph('LİSANS BİTİRME TEZİ')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.style.font.size = Pt(14)
    
    doc.add_paragraph('\n\n\n\n\n')
    
    # Student Info
    p = doc.add_paragraph(f'Hazırlayan\n{STUDENT_ID}')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.style.font.bold = True
    
    p = doc.add_paragraph(DEPT)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    doc.add_paragraph('\n\n\n')
    
    # Date
    p = doc.add_paragraph('KOCAELİ, Ocak 2026')
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.style.font.bold = True
    
    doc.add_page_break()

def clean_text_content(text):
    """
    Cleans markdown artifacts and unwanted Unicode symbols from text.
    """
    # Remove markdown link syntax [Text](#anchor) -> Text, or [Text](url) -> Text
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    
    # Remove bold markers **
    text = text.replace('**', '')
    
    # Remove italic markers * 
    text = text.replace('*', '')
    
    # Remove hashtags which might be leftover markdown
    text = text.replace('#', '')

    # Remove unwanted symbols
    text = re.sub(r'[✓✔☑]', '', text)
    
    return text.strip()

def add_formatted_cell(cell, text, is_header=False):
    cell._element.clear_content()
    p = cell.add_paragraph()
    
    # Clean the text for table cell
    clean_text = clean_text_content(text)
    
    run = p.add_run(clean_text)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    if is_header:
        run.font.bold = True
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    else:
        # Heuristic alignment
        if re.match(r'^[\d\.,%]+$', clean_text):
             p.alignment = WD_ALIGN_PARAGRAPH.CENTER 
        else:
             p.alignment = WD_ALIGN_PARAGRAPH.LEFT

def add_table_to_doc(doc, table_lines):
    if not table_lines:
        return

    # Parse rows
    rows = []
    for line in table_lines:
        # Split by pipe and strip whitespace
        parts = [p.strip() for p in line.split('|')]
        if len(parts) > 0 and parts[0] == '': parts = parts[1:]
        if len(parts) > 0 and parts[-1] == '': parts = parts[:-1]
        rows.append(parts)

    if len(rows) < 2: 
        return

    header = rows[0]
    separator_idx = 1
    if len(rows) > 1 and all('-' in cell for cell in rows[1]):
        separator_idx = 2
    
    data_rows = rows[separator_idx:]

    table = doc.add_table(rows=1, cols=len(header))
    table.style = 'Table Grid'
    
    # Set header
    hdr_cells = table.rows[0].cells
    for i, text in enumerate(header):
        if i < len(hdr_cells):
            add_formatted_cell(hdr_cells[i], text, is_header=True)

    # Add data rows
    for row_data in data_rows:
        row = table.add_row().cells
        for i, text in enumerate(row_data):
            if i < len(row):
                add_formatted_cell(row[i], text, is_header=False)
                
    doc.add_paragraph() 

def parse_markdown_and_add_to_doc(doc, md_path):
    if not os.path.exists(md_path):
        doc.add_paragraph("Error: Source markdown file not found.")
        return

    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    current_section = ""
    table_buffer = []
    skip_toc_section = False
    
    for line in lines:
        line = line.strip()
        
        # Check if we are inside a table
        if line.startswith('|'):
            table_buffer.append(line)
            continue
        elif table_buffer:
            add_table_to_doc(doc, table_buffer)
            table_buffer = []
        
        if not line:
            continue
            
        if line.startswith('# '): # Main Title - skip
            continue
            
        # Robust Header Detection using Regex
        # Matches ## Header, ### Header, ####Header, ##Header etc.
        header_match = re.match(r'^(#{2,6})\s*(.*)', line)
        if header_match:
            hashes = header_match.group(1)
            text = header_match.group(2)
            
            level = len(hashes)
            # Map level 1 (##) to Doc Heading 1, etc.
            # If level is 4 (####), it maps to Heading 3 usually in simple processing, 
            # or Heading 4 if we generated that style.
            # Let's cap at 3 because we only styled 1-3.
            
            doc_level = level - 1 # ## -> 1, ### -> 2, #### -> 3
            if doc_level > 3: doc_level = 3
            if doc_level < 1: doc_level = 1
            
            header_text = clean_text_content(text)
            
            # Check for TOC section to skip
            if "İÇİNDEKİLER" in header_text.upper():
                skip_toc_section = True
                continue
            else:
                if skip_toc_section:
                    skip_toc_section = False
            
            doc.add_heading(header_text, level=doc_level)
            
            if doc_level == 1:
                insert_images_for_section(doc, header_text)
            elif doc_level == 2:
                insert_images_for_subsection(doc, header_text)
                
            continue # Done with this line
        
        # If we are in TOC section and skipping
        if skip_toc_section:
            continue
            
        elif line.startswith('- ') or line.startswith('* '): # Bullet
            content = line[2:].strip()
            add_formatted_paragraph(doc, content, style='List Bullet')
            
        elif line.startswith('1. '): # Numbered
            content = line[3:].strip()
            add_formatted_paragraph(doc, content, style='List Number')
            
        elif line.startswith('$$'): 
            pass 
            
        else:
            if line == '---' or line == '___': 
                pass
            else:
                add_formatted_paragraph(doc, line)
    
    if table_buffer:
        add_table_to_doc(doc, table_buffer)

def insert_image(doc, filename, caption):
    full_path = os.path.join(FIGURES_DIR, filename)
    if os.path.exists(full_path):
        try:
            doc.add_picture(full_path, width=Inches(6.0))
            last_paragraph = doc.paragraphs[-1] 
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # Caption
            c = doc.add_paragraph(f'Şekil: {caption}')
            c.alignment = WD_ALIGN_PARAGRAPH.CENTER
            c.style.font.name = 'Times New Roman'
            c.style.font.size = Pt(12)
            doc.add_paragraph() # Spacer
        except Exception as e:
            print(f"Error adding image {filename}: {e}")

def insert_images_for_section(doc, header):
    # This logic wasn't working well because H1 sections in the file are just "BULGULAR", etc.
    # We will rely almost entirely on H2/H3 subsection matching.
    if "BULGULAR" in header:
        # Just a general placeholder or intro image if we had one
        pass

def insert_images_for_subsection(doc, header):
    # Consolidated image insertion logic here based on subsection titles
    if "Keşifsel Veri Analizi" in header:
        insert_image(doc, 'bolge_dagilimi.png', 'Bölgelere Göre İl Dağılımı')
        insert_image(doc, 'korelasyon_matrisi.png', 'Değişkenler Arası Korelasyon Matrisi (Heatmap)')
        insert_image(doc, 'dagilimlar.png', 'Önemli Değişkenlerin Dağılımları')
        insert_image(doc, 'pca_analizi.png', 'PCA Analizi ve 2D Görselleştirme')
        
    elif "Optimal Küme Sayısı" in header:
        insert_image(doc, 'elbow.png', 'Elbow (Dirsek) Yöntemi ile Optimal K')
        insert_image(doc, 'optimal_k_metrikleri.png', 'Silhouette ve Calinski-Harabasz Skorları')
        
    elif "K-Means Kümeleme" in header:
        # Generic K-Means or first intro
        pass
        
    elif "Küme Dağılımı" in header:
        insert_image(doc, 'kmeans_dagilim.png', 'K-Means Kümelerinin Harita Üzerinde Dağılımı')
    
    elif "Küme Profilleri" in header:
        insert_image(doc, 'kume_profilleri.png', 'Küme Profilleri (Ortalama Değerler)')
        insert_image(doc, 'kume_boxplot.png', 'Kümelerin Değişken Bazlı Dağılımları')
        insert_image(doc, 'kmeans_pca.png', 'K-Means Kümelerinin PCA Üzerinde Gösterimi')
        
    elif "Hiyerarşik Kümeleme" in header:
        insert_image(doc, 'dendrogram.png', 'Hiyerarşik Kümeleme Dendrogramı')
        
    elif "SEGE" in header and "Karşılaştırma" in header:
        insert_image(doc, 'sege_karsilastirma.png', 'Kümeler ve SEGE Kademeleri Karşılaştırması')

def main():
    doc = Document()
    setup_styles(doc)
    create_cover_page(doc)
    
    # Add TOC placeholder
    doc.add_heading('İÇİNDEKİLER', level=1)
    doc.add_paragraph('1. GİRİŞ........................................................................................... 1')
    doc.add_paragraph('2. LİTERATÜR TARAMASI................................................................... 3')
    doc.add_paragraph('3. METODOLOJİ................................................................................. 5')
    doc.add_paragraph('4. BULGULAR.................................................................................... 8')
    doc.add_paragraph('5. TARTIŞMA..................................................................................... 15')
    doc.add_paragraph('6. SONUÇ VE ÖNERİLER................................................................... 18')
    doc.add_page_break()
    
    parse_markdown_and_add_to_doc(doc, REPORT_MD_PATH)
    
    # Save
    if not os.path.exists(os.path.dirname(OUTPUT_DOCX_PATH)):
        os.makedirs(os.path.dirname(OUTPUT_DOCX_PATH))
        
    doc.save(OUTPUT_DOCX_PATH)
    print(f"Report saved to {OUTPUT_DOCX_PATH}")

if __name__ == "__main__":
    main()
